import io
import re
import xml.etree.ElementTree as ET
import zipfile
import concurrent.futures


try:
    import pypdf
    print("DEBUG: pypdf successfully imported inside parser_service!")
except ImportError:
    pypdf = None
    print("DEBUG: pypdf import failed! Using fallback text extractor.")

try:
    import docx
    print("DEBUG: docx successfully imported inside parser_service!")
except ImportError:
    docx = None
    print("DEBUG: docx import failed! Using fallback text extractor.")


def extract_text_with_timeout(func, *args, timeout=4.0):
    """
    Runs a parsing function in a separate thread with a strict timeout.
    Prevents malformed PDF infinite loops from freezing the FastAPI event loop.
    """
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(func, *args)
        try:
            return future.result(timeout=timeout)
        except concurrent.futures.TimeoutError:
            print(f"Parsing execution timed out after {timeout} seconds. Safe fallback triggered.")
            raise TimeoutError("Parsing timed out")


def extract_text_from_pdf_bytes(file_bytes: bytes) -> tuple[str, bool]:
    """
    Extracts text from a PDF file using pypdf.
    Returns (extracted_text, is_scanned_pdf).
    """
    def do_extract():
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        text_pieces = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_pieces.append(t)
        return "\n".join(text_pieces)

    if not pypdf:
        text = extract_text_pdf_fallback(file_bytes)
        is_scanned = len(text.strip()) < 100
        return text, is_scanned

    try:
        # Wrap pypdf in a strict 4.0s thread timeout to block infinite loops
        full_text = extract_text_with_timeout(do_extract, timeout=4.0)
        is_scanned = len(full_text.strip()) < 100
        return full_text, is_scanned
    except Exception as e:
        print(f"pypdf extraction failed or timed out: {e}")
        text = extract_text_pdf_fallback(file_bytes)
        return text, len(text.strip()) < 100


def extract_text_pdf_fallback(file_bytes: bytes) -> str:
    """Fallback pure-python regex PDF stream text extractor."""
    try:
        content = file_bytes.decode('latin1', errors='ignore')
        matches = re.findall(r'\(([^)]+)\)\s*(?:Tj|TJ)', content)
        if not matches:
            matches = re.findall(r'\(([^)]+)\)', content)
            matches = [m for m in matches if len(m) > 3 and not m.startswith('/') and not m.replace(' ', '').isdigit()]
        text = " ".join(matches)
        return text.replace('\\(', '(').replace('\\)', ')').replace('\\\\', '\\')
    except Exception:
        return ""


def extract_text_from_docx_bytes(file_bytes: bytes) -> tuple[str, bool]:
    """
    Extracts text from a DOCX file using python-docx.
    Returns (extracted_text, has_tables).
    """
    def do_extract():
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text_pieces = []
        
        # Read paragraphs
        for p in doc.paragraphs:
            if p.text:
                text_pieces.append(p.text)
                
        # Read tables
        has_tables = len(doc.tables) > 0
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)
                if row_text:
                    text_pieces.append(" ".join(row_text))
                    
        return "\n".join(text_pieces), has_tables

    if not docx:
        text = extract_text_docx_fallback(file_bytes)
        has_tables = b"<w:tbl" in file_bytes
        return text, has_tables

    try:
        full_text, has_tables = extract_text_with_timeout(do_extract, timeout=4.0)
        return full_text, has_tables
    except Exception as e:
        print(f"python-docx extraction failed or timed out: {e}")
        text = extract_text_docx_fallback(file_bytes)
        return text, b"<w:tbl" in file_bytes


def extract_text_docx_fallback(file_bytes: bytes) -> str:
    """Fallback pure-python DOCX text extractor."""
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            xml_content = z.read("word/document.xml")
            root = ET.fromstring(xml_content)
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            texts = [node.text for node in root.findall('.//w:t', namespaces) if node.text]
            return " ".join(texts)
    except Exception:
        return ""


def extract_file_content(filename: str, file_bytes: bytes) -> dict:
    """
    Ingests file name and file bytes, identifies format, extracts text,
    and runs specific diagnostics for table presence and image checks.
    """
    ext = filename.split('.')[-1].lower() if '.' in filename else ''
    text = ""
    has_tables = False
    is_scanned = False
    status = "Successfully parsed"

    try:
        if ext == 'pdf':
            text, is_scanned = extract_text_from_pdf_bytes(file_bytes)
            if is_scanned:
                status = "Scanned image detected (Workday/Greenhouse alert)"
        elif ext in ['docx', 'doc']:
            text, has_tables = extract_text_from_docx_bytes(file_bytes)
            if has_tables:
                status = "Tables found (potential Workday formatting alert)"
        elif ext in ['txt', 'rtf']:
            text = file_bytes.decode('utf-8', errors='ignore')
        else:
            text = file_bytes.decode('utf-8', errors='ignore')
            status = "Unknown extension: falling back to plain text parsing"
    except Exception as e:
        status = f"Extraction error: {str(e)}"

    return {
        "text": text,
        "has_tables": has_tables,
        "is_scanned": is_scanned,
        "status": status,
        "length": len(text)
    }
